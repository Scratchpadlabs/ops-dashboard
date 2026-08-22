#!/usr/bin/env python3
"""
ClarifiEd material intake — School Material Import Cloud Functions.

Ports the extraction/validation logic from the `extract.py` prototype into a
callable Cloud Function (`process_import`) that writes results straight into
Firestore `staging_imports/{jobId}` (job doc + `rows/{n}` subcollection),
instead of CSV files, so the review UI can stream progress and 1000+ row
imports don't have to round-trip through the browser. `commit_import` is the
second half — it takes the already-classified plan the review UI built
client-side (src/composables/useImport.js's buildCommitPlan, still plain
Firestore reads, unchanged) and performs the actual writes into
`schools/{schoolId}/...` server-side.

Both are Firebase `on_call` callables, not raw HTTP functions: the callable
protocol handles CORS/preflight itself (no manual OPTIONS branch or
Access-Control-* headers needed) and gives each function a verified,
decoded Firebase Auth ID token at `req.auth` — no custom X-Api-Key header,
which is what raw `fetch()` + a hardcoded key both required and got wrong
(a plain HTTP Cloud Function that isn't itself invokable anonymously fails
IAM auth before your CORS headers ever get a chance to be written, which the
browser reports as a CORS error).

Deploy (per function, same source directory) — this repo deploys via plain
`gcloud functions deploy`, not `firebase deploy --only functions`, so the
`@https_fn.on_call(...)` decorator's memory/timeout_sec/max_instances/secrets
arguments are NOT read at deploy time (that introspection only happens
through the Firebase CLI's own build pipeline). They still matter for the
callable *behavior* (CORS/auth), but the actual Cloud Run resource limits
come from these gcloud flags and must be kept in sync with the decorator:
  gcloud functions deploy process_import \
    --gen2 --runtime python312 --region asia-south1 \
    --source . --entry-point process_import \
    --trigger-http --allow-unauthenticated \
    --memory 1024MB --timeout 540s --max-instances 3 --project clarified-1501 \
    --set-secrets OPENAI_API_KEY=OPENAI_API_KEY:latest \
    --set-env-vars MODEL=gpt-4o
  # (bind ANTHROPIC_API_KEY=ANTHROPIC_API_KEY:latest instead — and drop the
  # MODEL override, or set it to a Claude model — to extract from .pdf
  # uploads, which the OpenAI path can't handle; see call_openai_compatible)

  gcloud functions deploy commit_import \
    --gen2 --runtime python312 --region asia-south1 \
    --source . --entry-point commit_import \
    --trigger-http --allow-unauthenticated \
    --memory 512MB --timeout 120s --max-instances 3 --project clarified-1501

`--allow-unauthenticated` is still required at the IAM layer even for
callables — that only lets the request reach the function; the function
itself then checks `req.auth` (the Firebase ID token) and the ops-admin
allowlist before doing anything. See functions/DEPLOY.md for the Secret
Manager setup.

The function's runtime service account needs `roles/datastore.user` (Firestore
read/write) and `roles/storage.objectViewer` (read uploaded source files) on
the clarified-1501 project — grant those once via the Cloud Console/gcloud,
they aren't set up by this deploy command.
"""
import base64, json, os, re, statistics
from datetime import datetime, date
from pathlib import PurePosixPath

import requests

import firebase_admin
from firebase_admin import firestore, storage
from firebase_functions import https_fn, options

from normalize import (
    ROMAN_TO_NUM, normalize_grade, normalize_section, canonicalize,
    extract_grade_tokens, clean_name, clean_email, EMAIL_DOMAIN_FIXES, _EMAIL_RE,
    FUZZY_THRESHOLD, fuzzy_best_match, match_value,
    STUDENT_HEADER_ALIASES, STUDENT_SCHEMA_KEYS, STUDENT_REQUIRED_FIELD,
    REVIEW_ONLY_STUDENT_KEYS,
)
import education_kb as kb
from tabular_parser import parse_tabular_file
# ONE copy, imported — functions/shared/ is staged into the deploy tree by
# tools/deploy_function.sh (see that script and functions/shared/__init__.py).
from shared.school_schema import (
    validate_doc, format_errors, validate_current_class_id, coerce_wire_payload,
)
from shared.class_resolver import class_id_key

# Must run at module load, not lazily inside a handler: the on_call framework
# verifies the caller's Firebase Auth ID token BEFORE our function body ever
# runs, and that verification itself needs the Admin SDK already initialized
# to decode the token. A lazy _ensure_firebase() called from inside the
# handler is too late — on a cold start with a real (non-empty) ID token,
# verification crashes with "The default Firebase app does not exist" before
# our code, including our own req.auth check, ever executes.
# storageBucket must be explicit — unlike some Admin SDKs, initialize_app()
# doesn't infer it, and storage.bucket() then fails with "Storage bucket
# name not specified" (matches src/firebase/config.js's storageBucket).
firebase_admin.initialize_app(options={"storageBucket": "clarified-1501.appspot.com"})

# No hardcoded default here — call_anthropic/call_openai_compatible each pick
# their own provider-appropriate default when MODEL isn't set (see extract_file
# for the provider switch: OpenAI when OPENAI_API_KEY is set, else Anthropic).
MODEL = os.environ.get("MODEL")

# Mirrors src/config/opsAdmins.js — keep in sync. Server-side is the
# authoritative check; the frontend's isOpsAdmin() is only a UI gate.
OPS_ADMIN_EMAILS = {"sid@ops.clarified.in", "angel@ops.clarified.in"}


def _require_ops_admin(req: https_fn.CallableRequest) -> str:
    """Verifies the callable's Firebase Auth token and the ops-admin
    allowlist server-side. Returns the caller's email for use in
    created_by/updated_by fields."""
    if req.auth is None:
        raise https_fn.HttpsError(
            https_fn.FunctionsErrorCode.UNAUTHENTICATED, "Sign in required.")
    email = str((req.auth.token or {}).get("email") or "").strip().lower()
    if email not in OPS_ADMIN_EMAILS:
        raise https_fn.HttpsError(
            https_fn.FunctionsErrorCode.PERMISSION_DENIED,
            "Not authorized for School Material Import.")
    return email

# ---------------------------------------------------------------- schemas ---
# Kept as-is from extract.py — the prompts already forbid Aadhaar/SSSM/caste/
# religion/address (golden rule 3); since stored rows are built by reading
# ONLY these declared keys off the LLM's JSON, any hallucinated extra field
# (banned or not) is dropped automatically before it ever reaches Firestore.
SCHEMAS = {
    "students": {
        # Kept identical to normalize.STUDENT_SCHEMA_KEYS — the deterministic
        # tabular_parser path and this LLM-extraction path must agree
        # byte-for-byte on the row shape, since both feed the same
        # clean_students_rows/validate_students below.
        "row": STUDENT_SCHEMA_KEYS,
        "hints": (
            "grade: Nursery/LKG/UKG or 1..12 as plain numbers (convert roman numerals). "
            "section: single letter or empty. dob: YYYY-MM-DD; Excel serial numbers are "
            "days since 1899-12-30. Strip honorifics (MAS./MISS/MRS./MR.) from names and "
            "use them to infer gender. contact: first valid 10-digit number only. adm_no "
            "is the admission/GR register number, distinct from sr_no (serial number). "
            "DO NOT extract Aadhaar numbers, SSSM ids, caste/category, religion, or "
            "addresses even if present — omit them entirely."
        ),
        "required": ["grade", "student_name"],
    },
    "teachers": {
        "row": ["teacher_name", "email", "class_teacher_of", "subject", "grade", "section"],
        "hints": (
            "Output ONE ROW PER (teacher, subject, class-section). Expand ranges: "
            "'VI - A,B,C' becomes three rows grade=6 sections A,B,C; '3 to 8' expands "
            "each grade with empty section. In matrix layouts the column header is the "
            "subject — read column alignment carefully. If a cell's column is visually "
            "ambiguous, still output the row but append '?' to the subject."
        ),
        "required": ["teacher_name", "subject", "grade"],
    },
    "subjects": {
        "row": ["stream", "grade_band", "subject", "area"],
        "hints": "area is 'Scholastic' or 'Co-Scholastic'. stream e.g. CBSE / International / SA; empty if not stated.",
        "required": ["subject"],
    },
    "assessments": {
        # `subject` is optional and usually absent: a blueprint row describes an
        # exam for a whole grade band, and the browser fans it out over that
        # band's subjects (see src/utils/assessmentImport.js). It is read when
        # the file DOES name one subject per row, so those rows land on that
        # subject instead of the whole band.
        "row": ["stream", "grade_band", "subject", "assessment", "date_start",
                "date_end", "instructional_days", "syllabus_covered",
                "exam_syllabus", "max_written", "activity_weight", "total",
                "duration"],
        "hints": ("Dates as YYYY-MM-DD. Keep composite marks like '80 (40+40)' "
                  "verbatim in max_written, and put the overall paper total in "
                  "`total` when the file states one. subject: only when the row "
                  "is for ONE named subject; leave empty for a grade-wide exam "
                  "blueprint."),
        "required": ["assessment"],
    },
}

# Fields the extractor must never pull out of a source file (golden rule 3).
#
# "aadhaar" was removed from this set on 2026-08-04 by an explicit decision to
# persist it on the student document. NOTE the consequence, which the ban was
# protecting against: firestore.rules grants
#   match /schools/{schoolId}/{collection}/{docId} { allow read: if isAuthenticated(); }
# so every student document is readable by ANY signed-in user of the teacher
# and student apps — the write path being ops-admin-only does not narrow that.
# Restricting who can READ Aadhaar needs a rules change, not a code change.
BANNED_KEYS = {"sssm_id", "sssm", "caste", "category", "religion", "address"}

# Header-level twin of the above, for the DETERMINISTIC path's `extras` — the
# columns the alias dictionary has no field for, carried through so a school
# can get its own data onto its own students.
#
# Shorter than BANNED_KEYS on purpose: "address" and "category" were allowed
# on 2026-08-20 by explicit decision, the same way "aadhaar" was on 2026-08-04,
# and the same consequence applies — firestore.rules lets any signed-in user of
# the teacher and student apps read every student document, so an allowed field
# is a readable field. Narrowing that needs a rules change, not a code change.
# Matched against canonicalize(header), so "Caste Category" and "caste-category"
# are the same header.
BANNED_EXTRA_TOKENS = ("sssm", "caste", "religion")


def is_banned_extra(header):
    """Whether a column header names something golden rule 3 forbids storing.

    Substring, not equality: the header is a school's own free text, so the
    ban has to catch "Caste", "Caste Category", "Student Caste (SC/ST)" and
    "SSSM ID" alike. Over-matching here costs a column an operator can add by
    hand; under-matching writes a caste record into a document every signed-in
    app user can read.
    """
    key = canonicalize(header)
    return any(token in key for token in BANNED_EXTRA_TOKENS)

# ------------------------------------------------------------- preprocess ---
def xlsx_to_tsv(raw: bytes) -> str:
    import io
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    out = []
    for name in wb.sheetnames:
        out.append(f"## Sheet: {name}")
        for row in wb[name].iter_rows(values_only=True):
            if any(v is not None and str(v).strip() for v in row):
                out.append("\t".join("" if v is None else str(v) for v in row))
    return "\n".join(out)

def docx_to_text(raw: bytes) -> str:
    import io
    import docx
    d = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for r in t.rows:
            parts.append("\t".join(c.text.strip() for c in r.cells))
        parts.append("")
    return "\n".join(parts)

def content_blocks(filename: str, raw: bytes):
    """Returns Anthropic-style content blocks for the file."""
    ext = PurePosixPath(filename).suffix.lower()
    if ext in (".xlsx", ".xlsm"):
        return [{"type": "text", "text": xlsx_to_tsv(raw)}]
    if ext == ".docx":
        return [{"type": "text", "text": docx_to_text(raw)}]
    if ext == ".pdf":
        data = base64.b64encode(raw).decode()
        return [{"type": "document",
                 "source": {"type": "base64", "media_type": "application/pdf", "data": data}}]
    if ext in (".png", ".jpg", ".jpeg", ".webp"):
        media = "image/png" if ext == ".png" else "image/jpeg"
        data = base64.b64encode(raw).decode()
        return [{"type": "image",
                 "source": {"type": "base64", "media_type": media, "data": data}}]
    return [{"type": "text", "text": raw.decode(errors="replace")}]

# -------------------------------------------------------------- llm calls ---
def build_prompt(entity: str) -> str:
    s = SCHEMAS[entity]
    return (
        "You are a data-extraction engine for an Indian school report-card platform. "
        f"Extract ALL {entity} records from the attached material.\n"
        f"Rules: {s['hints']}\n"
        "Return ONLY a JSON array of objects (no markdown, no commentary) with exactly "
        f"these keys: {json.dumps(s['row'])}. Use empty string for unknown values. "
        "Never invent data not present in the source."
    )

def call_anthropic(blocks, prompt):
    r = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={"x-api-key": os.environ["ANTHROPIC_API_KEY"],
                 "anthropic-version": "2023-06-01"},
        json={"model": MODEL or "claude-sonnet-4-6", "max_tokens": 16000,
              "messages": [{"role": "user",
                            "content": blocks + [{"type": "text", "text": prompt}]}]},
        timeout=300)
    r.raise_for_status()
    return "".join(b.get("text", "") for b in r.json()["content"])

def call_openai_compatible(blocks, prompt):
    # Defaults to the real OpenAI API; override OPENAI_BASE_URL to point at a
    # self-hosted OpenAI-compatible endpoint instead.
    content = []
    for b in blocks:
        if b["type"] == "text":
            content.append({"type": "text", "text": b["text"]})
        elif b["type"] == "image":
            url = f'data:{b["source"]["media_type"]};base64,{b["source"]["data"]}'
            content.append({"type": "image_url", "image_url": {"url": url}})
        else:
            # OpenAI's Chat Completions API has no raw-PDF content block (unlike
            # Anthropic's Messages API) — .pdf uploads aren't extractable on this
            # path without rasterizing pages to images first.
            raise RuntimeError("PDF input isn't supported via the OpenAI extraction path (rasterize pages to PNG first, or set ANTHROPIC_API_KEY instead).")
    content.append({"type": "text", "text": prompt})
    base_url = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")
    r = requests.post(
        base_url.rstrip("/") + "/chat/completions",
        headers={"Authorization": "Bearer " + os.environ.get("OPENAI_API_KEY", "x")},
        json={"model": MODEL or "gpt-4o",
              "messages": [{"role": "user", "content": content}],
              "max_tokens": 16000},
        timeout=300)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

def extract_file(entity, filename, raw):
    blocks = content_blocks(filename, raw)
    prompt = build_prompt(entity)
    raw_text = (call_openai_compatible if os.environ.get("OPENAI_API_KEY")
                else call_anthropic)(blocks, prompt)
    raw_text = re.sub(r"^```(json)?|```$", "", raw_text.strip(), flags=re.M).strip()
    rows = json.loads(raw_text)
    schema_keys = SCHEMAS[entity]["row"]
    # Only ever keep declared schema keys — drops any hallucinated PII field
    # regardless of prompt compliance (golden rule 3, belt-and-braces).
    return [{k: str(r.get(k, "") or "") for k in schema_keys} for r in rows]

# --------------------------------------------------------------- cleaning ---
# Real school sheets are messy — typo email domains, "Garde 7,10", stray
# whitespace, "Eng" vs "English" — and without this stage nearly every row
# gets flagged/skipped even when the data is trivially recoverable. This runs
# on every parsed row BEFORE validate_*/flag generation below. Every auto-fix
# is recorded as {field, original, fixed, rule} so review stays lossless:
# originals are always visible, nothing here writes to the live DB directly,
# and anything uncertain becomes a *suggestion* (shown for one-click accept
# in review) or a hard flag — never a silent guess.
#
# name/email/grade/section normalizers now live in normalize.py (imported
# above) — shared verbatim with tabular_parser.py's deterministic path, per
# the hardening task's "one shared module" requirement.

def load_kb_entries(db):
    """Learned overlay for the education knowledge base (top-level
    `kb_entries`, not scoped to any school) — the same self-learning pattern
    as import_aliases, but for MEANING rather than spelling: {canonicalized
    value: (type, canonical)}.

    Only ever written by a human confirming a classification in the UI (see
    the classify_value callable below and src/composables/useEducationKB.js).
    A cached type of "other" is as useful as a positive answer — it stops the
    LLM being asked about that value ever again."""
    try:
        return kb.build_overlay([d.to_dict() or {} for d in db.collection("kb_entries").stream()])
    except Exception as e:
        # A missing/unreadable overlay must never take down an import — the
        # seeded KB alone still classifies the overwhelming majority.
        print(f"load_kb_entries: falling back to seed-only KB: {e}")
        return {}


def kb_match(raw, kind, valid_canon_to_display, aliases, kb_overlay):
    """match_value(), plus the knowledge base as a second chance.

    The school's own configured values win first (a school that calls it
    'Eng' in its live Subjects list should keep seeing 'Eng'). Only when that
    finds nothing does the KB get consulted, to turn the raw value into its
    canonical form ('Eng' -> 'English') and retry the match — which is what
    lets an import land on a school's configured 'English' with no learned
    alias and no fuzzy guesswork.

    Returns match_value's own (status, display, extra) triple, so callers
    treat a KB-resolved hit exactly like any other auto-fix."""
    status, display, extra = match_value(raw, kind, valid_canon_to_display, aliases)
    if status == "auto":
        return status, display, extra

    expect = kb.SUBJECT if kind == "subject" else kb.SECTION
    result = kb.classify(raw, expect=expect, overlay=kb_overlay)
    canonical = result["canonical"]
    if not canonical or result["confidence"] < kb.HIGH_CONFIDENCE:
        return status, display, extra
    if result["type"] not in (kb.SUBJECT, kb.COSCHOLASTIC, kb.SECTION):
        return status, display, extra

    retry_status, retry_display, _ = match_value(canonical, kind, valid_canon_to_display, aliases)
    if retry_status == "auto":
        return "auto", retry_display, "kb_canonical_match"
    # The KB knows what the value means but this school hasn't configured it —
    # keep the weaker original outcome rather than inventing a value the
    # school's setup can't accept. Structure inference (Part C) is where a
    # missing subject gets proposed for real.
    return status, display, extra


def load_import_aliases(db):
    """Global alias map (top-level `import_aliases` collection, not scoped to
    any one school) — {(type, canonical_from): display_to}. An exact hit here
    is an auto-fix, applied the same way for every school. Sid accepting a
    suggestion (or manually mapping a value) in review writes back here
    client-side, so later imports for ANY school auto-apply it — see
    src/composables/useImport.js's acceptSuggestion."""
    aliases = {}
    for d in db.collection("import_aliases").stream():
        a = d.to_dict() or {}
        t, frm, to = a.get("type"), a.get("from"), a.get("to")
        if t and frm and to:
            aliases[(t, frm)] = to
    return aliases


def clean_students_rows(rows, cfg, provenance=None):
    """Stage 1 + 2 for students: name cleanup, tolerant grade parsing (first
    token only — a student has one grade; multiple tokens is itself flagged
    as ambiguous, not expanded), and section canonicalize/alias/fuzzy against
    that grade's configured sections. Returns a list of
    {data, fixes, suggestions, flags, provenance} — `flags` here only covers
    what validate_students below has no way to catch on its own (an ambiguous
    multi-grade cell); everything else it flags is appended afterward, once
    run over this (now cleaner) `data`. `provenance` (optional, same length
    as `rows`) carries {source_file, sheet, row} through unchanged — students
    never expand 1 input row into >1 output row, so the alignment is exact."""
    out = []
    for i, r in enumerate(rows):
        d = dict(r)
        fixes, suggestions, flags = [], [], []

        for field in ("student_name", "mother_name", "father_name"):
            v, fix = clean_name(d.get(field, ""))
            d[field] = v
            if fix:
                fixes.append({**fix, "field": field})

        grade_raw = d.get("grade", "")
        tokens = extract_grade_tokens(grade_raw)
        if tokens:
            grade = tokens[0]
            if grade != grade_raw.strip():
                fixes.append({"field": "grade", "original": grade_raw, "fixed": grade, "rule": "grade_parsed"})
            d["grade"] = grade
            if len(tokens) > 1:
                flags.append({"field": "grade", "message": f"multiple grades found in '{grade_raw}' — using '{grade}', please verify", "severity": "warning"})
        else:
            grade = ""

        section_raw = d.get("section", "")
        if section_raw.strip():
            candidates = cfg["sections_by_grade"].get(grade, {})
            status, display, extra = kb_match(section_raw, "class", candidates, cfg["aliases"], cfg["kb_overlay"])
            if status == "auto":
                if display != section_raw.strip():
                    fixes.append({"field": "section", "original": section_raw, "fixed": display, "rule": extra})
                d["section"] = display
            elif status == "suggestion":
                suggestions.append({"field": "section", "original": section_raw, "suggested": display, "similarity": round(extra, 2)})

        prov = provenance[i] if provenance else None
        out.append({"data": d, "fixes": fixes, "suggestions": suggestions, "flags": flags, "provenance": prov})
    return out


def clean_teachers_rows(rows, cfg, provenance=None):
    """Stage 1 + 2 for teachers. Grade is the one field that can legitimately
    EXPAND a row: a multi-grade cell ('Garde 7,10') means one teacher-subject
    assignment spanning several grades, which the LLM is supposed to expand
    into separate rows itself but doesn't always manage — this is the
    deterministic fallback. Section/subject are then canonicalized per the
    row's own (possibly newly-split) grade. Returns a list of
    {data, fixes, suggestions, flags, provenance}, one per OUTPUT row (>=
    len(rows)); `flags` is always empty here — a multi-grade cell is
    expected/valid for teachers, not ambiguous, so there's nothing cleaning
    itself needs to flag (unlike students) — validate_teachers appends
    everything else. `provenance` (optional, same length as `rows`) is
    duplicated onto every row an expansion produces from the same source."""
    out = []
    for i, r in enumerate(rows):
        prov = provenance[i] if provenance else None
        base_fixes = []

        v, fix = clean_name(r.get("teacher_name", ""))
        if fix:
            base_fixes.append({**fix, "field": "teacher_name"})

        email_raw = r.get("email", "")
        # Validity itself isn't checked here — validate_teachers below raises
        # the hard flag for a missing "@" or still-invalid address, same as
        # every other unresolved field, once run over this cleaned data.
        email_v, email_fix, _ = clean_email(email_raw)
        if email_fix:
            base_fixes.append(email_fix)

        grade_raw = r.get("grade", "")
        tokens = extract_grade_tokens(grade_raw)
        grades = tokens if tokens else [""]

        for grade in grades:
            d = dict(r)
            d["teacher_name"] = v
            d["email"] = email_v
            fixes = list(base_fixes)
            suggestions = []

            if grade:
                d["grade"] = grade
                if grade != grade_raw.strip() or len(tokens) > 1:
                    rule = "grade_parsed_expanded" if len(tokens) > 1 else "grade_parsed"
                    fixes.append({"field": "grade", "original": grade_raw, "fixed": grade, "rule": rule})

            section_raw = d.get("section", "")
            if section_raw.strip():
                candidates = cfg["sections_by_grade"].get(grade, {})
                status, display, extra = kb_match(section_raw, "class", candidates, cfg["aliases"], cfg["kb_overlay"])
                if status == "auto":
                    if display != section_raw.strip():
                        fixes.append({"field": "section", "original": section_raw, "fixed": display, "rule": extra})
                    d["section"] = display
                elif status == "suggestion":
                    suggestions.append({"field": "section", "original": section_raw, "suggested": display, "similarity": round(extra, 2)})

            subject_raw = d.get("subject", "")
            subject_ambiguous = subject_raw.endswith("?")
            subject_clean = subject_raw.rstrip("?").strip()
            if subject_clean:
                candidates = cfg["subjects_by_grade_canon"].get(grade, {})
                status, display, extra = kb_match(subject_clean, "subject", candidates, cfg["aliases"], cfg["kb_overlay"])
                if status == "auto":
                    fixed_value = display + ("?" if subject_ambiguous else "")
                    if fixed_value != subject_raw:
                        fixes.append({"field": "subject", "original": subject_raw, "fixed": fixed_value, "rule": extra})
                    d["subject"] = fixed_value
                elif status == "suggestion":
                    suggestions.append({"field": "subject", "original": subject_raw, "suggested": display, "similarity": round(extra, 2)})

            out.append({"data": d, "fixes": fixes, "suggestions": suggestions, "flags": [], "provenance": prov})
    return out

def clean_subjects_rows(rows, cfg, provenance=None):
    """Stage 1 + 2 for subjects. The one entity where the knowledge base does
    the whole job on its own: it decides whether each row is a scholastic
    subject or a co-scholastic area (the `area` column schools leave blank
    more often than not) and normalizes the name to its canonical form.

    Both are auto-fixes only at HIGH confidence. A medium-confidence name
    becomes a suggestion for review; an unknown one is left exactly as typed
    with no `area` invented — an unrecognized activity name is precisely what
    the classify_value LLM fallback and a human confirmation exist for, and
    guessing 'Scholastic' here would quietly mis-file it."""
    out = []
    for i, r in enumerate(rows):
        d = dict(r)
        fixes, suggestions, flags = [], [], []

        name_raw = (d.get("subject") or "").strip()
        if name_raw:
            result = kb.classify(name_raw, overlay=cfg["kb_overlay"])
            canonical = result["canonical"]
            is_entity = result["type"] in (kb.SUBJECT, kb.COSCHOLASTIC)

            if is_entity and result["confidence"] >= kb.HIGH_CONFIDENCE:
                if canonical != name_raw:
                    fixes.append({"field": "subject", "original": name_raw, "fixed": canonical,
                                   "rule": f"kb_{result['source']}"})
                    d["subject"] = canonical
                area = "Scholastic" if result["type"] == kb.SUBJECT else "Co-Scholastic"
                if not (d.get("area") or "").strip():
                    fixes.append({"field": "area", "original": "", "fixed": area, "rule": "kb_area"})
                    d["area"] = area
            elif is_entity and result["confidence"] >= kb.MEDIUM_CONFIDENCE:
                suggestions.append({"field": "subject", "original": name_raw, "suggested": canonical,
                                     "similarity": round(result["confidence"], 2)})
            elif not (d.get("area") or "").strip():
                flags.append({"field": "area", "severity": "warning",
                               "message": f"'{name_raw}' is not in the education knowledge base — "
                                          "set Scholastic/Co-Scholastic manually, or classify it in School Setup"})

        prov = provenance[i] if provenance else None
        out.append({"data": d, "fixes": fixes, "suggestions": suggestions, "flags": flags, "provenance": prov})
    return out


# -------------------------------------------------------------- validate ----
def parse_dob(d):
    d = (d or "").strip()
    if not d or not re.match(r"^\d{4}-\d{2}-\d{2}$", d):
        return None
    try:
        return datetime.strptime(d, "%Y-%m-%d").date()
    except ValueError:
        return None

def _flag(field, message, severity="warning"):
    """Structured flag: {field, message, severity}. `field` is None for
    row-level checks not tied to one cell (duplicate name, DOB-vs-median) —
    lets the review UI clear only the flags a specific accepted suggestion
    resolves, instead of guessing from message text. `severity` is "error"
    only for the handful of checks that mean the row can't be committed
    as-is (missing the one required field, invalid email) — everything else
    is "warning": recoverable, kept, surfaced for human review."""
    return {"field": field, "message": message, "severity": severity}

def _already_flagged(existing, field):
    """True if the parser already reported something about this cell.

    tabular_parser blanks a cell it cannot read and reports why
    ("unparseable date of birth: '16 Jan 2024'"). Without this check the
    row-level pass then sees the blank and adds "missing date of birth" —
    two messages, one failure, and the second one is actively wrong: the
    file DID have a date. Same collision applies to gender and contact.
    """
    return any(f.get("field") == field for f in (existing or []))


def resolve_by_doc_id(class_by_doc_id, raw_grade, raw_section):
    """The class a row names by its DOCUMENT id, or None.

    A school's export can carry the class as ONE complete id instead of a
    Class + Section pair — Hillgreen's 2026-27 export dropped its Section
    column and put "Play_Group_A" / "8_KALAM" straight into Class. Because
    STUDENT_HEADER_ALIASES maps "class" onto grade, the (grade, section)
    lookup cannot see those, and every row was told its grade was "not
    configured for this school" when the class was configured all along.

    Only ever consulted AFTER the (grade, section) lookup misses, and an id
    that is not configured still returns None — this cannot invent a class.

    Twin of resolveClassId's fallback rungs in src/utils/classLookup.js: the
    commit planner in the browser and this validator have to agree about the
    same row, or the review screen warns about what the commit then accepts.
    """
    if not class_by_doc_id:
        return None
    gkey = class_id_key(raw_grade)
    if gkey and gkey in class_by_doc_id:
        return class_by_doc_id[gkey]
    skey = class_id_key(raw_section)
    if gkey and skey and f"{gkey}_{skey}" in class_by_doc_id:
        return class_by_doc_id[f"{gkey}_{skey}"]
    return None


UNRECOGNIZED_GRADE_PREFIX = "unrecognized grade value:"


def drop_unrecognized_grade_flags(cleaned, class_by_doc_id):
    """Retract the parser's "unrecognized grade value" once the school config
    shows the value was a class id all along.

    tabular_parser reads each cell with no knowledge of the school, so a
    Class cell holding a whole class id ("8_KALAM") is correctly reported as
    not being a grade. Only here, after load_school_config, can that be
    resolved — and on Hillgreen's 2026-27 export it was every single row:
    1621 warnings about 52 classes the school has configured.

    A row whose class does NOT resolve keeps its warning, which is the whole
    point: the report should be short enough that the real one is visible.
    """
    for c in cleaned:
        d = c.get("data") or {}
        if not resolve_by_doc_id(class_by_doc_id, d.get("grade"), d.get("section")):
            continue
        c["flags"] = [f for f in c["flags"]
                      if not (f.get("field") == "grade"
                              and str(f.get("message", "")).startswith(UNRECOGNIZED_GRADE_PREFIX))]
    return cleaned


def validate_students(rows, class_lookup, sections_by_grade, existing_flags=None,
                      class_by_doc_id=None):
    flags_by_row = [[] for _ in rows]
    seen_names = {}
    dobs_by_grade = {}
    parsed_dobs = [None] * len(rows)

    for i, r in enumerate(rows):
        if not r.get("student_name", "").strip():
            flags_by_row[i].append(_flag("student_name", "missing student name", severity="error"))
        if not r.get("grade", "").strip():
            flags_by_row[i].append(_flag("grade", "missing grade"))
        prior = (existing_flags or [None] * len(rows))[i] if existing_flags else []
        dob_raw = r.get("dob", "").strip()
        if not dob_raw:
            # Only "missing" when the cell really was empty — a cell the
            # parser rejected has already said so, in more useful words.
            if not _already_flagged(prior, "dob"):
                flags_by_row[i].append(_flag("dob", "missing date of birth"))
        else:
            d = parse_dob(dob_raw)
            if d is None:
                flags_by_row[i].append(_flag("dob", f"unparseable date of birth: '{dob_raw}'"))
            else:
                parsed_dobs[i] = d
                dobs_by_grade.setdefault(r.get("grade", ""), []).append(d)
        if not r.get("gender", "").strip() and not _already_flagged(prior, "gender"):
            flags_by_row[i].append(_flag("gender", "missing gender"))
        # A parent's number counts. father_mobile and mother_mobile are written
        # to the student document since 2026-08-20 (CARRIED_SOURCE_FIELDS in
        # src/schemas/studentMapping.js), so a row carrying one is not missing
        # a contact — it just does not fill the column this check used to read
        # alone. On Hillgreen's export that was 1472 of 1622 rows warned while
        # every single one carried a father's number; a warning on 90% of a
        # file is one nobody reads.
        has_contact = any(r.get(k, "").strip()
                          for k in ("contact", "father_mobile", "mother_mobile"))
        if not has_contact and not _already_flagged(prior, "contact"):
            flags_by_row[i].append(_flag("contact", "no contact number in any column"))

        key = (normalize_grade(r.get("grade")), normalize_section(r.get("section")),
               r.get("student_name", "").strip().lower())
        if key[2]:
            if key in seen_names:
                flags_by_row[i].append(_flag("student_name", "duplicate student name in this class-section"))
            seen_names[key] = True

        grade = normalize_grade(r.get("grade"))
        section = normalize_section(r.get("section"))
        gkey = (grade, section)
        if gkey not in class_lookup and not resolve_by_doc_id(class_by_doc_id, r.get("grade"), r.get("section")):
            if grade and grade not in sections_by_grade:
                flags_by_row[i].append(_flag("grade", f"grade '{grade}' not configured for this school"))
            elif section:
                flags_by_row[i].append(_flag("section", f"class '{r.get('section','').strip()}' not in school config for grade {grade}"))

    # DOB implausible for grade: >3 years off the grade's median.
    medians = {g: statistics.median([d.toordinal() for d in ds]) for g, ds in dobs_by_grade.items() if ds}
    for i, r in enumerate(rows):
        d = parsed_dobs[i]
        if d is None:
            continue
        med = medians.get(r.get("grade", ""))
        if med is None:
            continue
        years_off = abs(d.toordinal() - med) / 365.25
        if years_off > 3:
            flags_by_row[i].append(_flag("dob", f"date of birth implausible for grade ({years_off:.1f}y off class median)"))

    return flags_by_row

def validate_teachers(rows, class_lookup, sections_by_grade, subject_names_by_grade,
                      class_by_doc_id=None):
    flags_by_row = [[] for _ in rows]

    # Group by normalized teacher identity to find zero-assignment teachers.
    by_teacher = {}
    for i, r in enumerate(rows):
        tkey = (r.get("teacher_name", "").strip().lower(), r.get("email", "").strip().lower())
        by_teacher.setdefault(tkey, []).append(i)

    for i, r in enumerate(rows):
        if not r.get("teacher_name", "").strip():
            flags_by_row[i].append(_flag("teacher_name", "missing teacher name", severity="error"))

        email = r.get("email", "").strip()
        if email and not _EMAIL_RE.match(email):
            flags_by_row[i].append(_flag("email", f"invalid email: '{email}'", severity="error"))

        grade = normalize_grade(r.get("grade"))
        section = normalize_section(r.get("section"))
        gkey = (grade, section)
        if grade and gkey not in class_lookup and not resolve_by_doc_id(class_by_doc_id, r.get("grade"), r.get("section")):
            if grade not in sections_by_grade:
                flags_by_row[i].append(_flag("grade", f"grade '{grade}' not configured for this school"))
            elif section:
                flags_by_row[i].append(_flag("section", f"class '{r.get('section','').strip()}' not in school config for grade {grade}"))

        subject = r.get("subject", "").strip().rstrip("?")
        if subject:
            allowed = subject_names_by_grade.get(grade)
            if allowed is not None and subject.lower() not in allowed:
                flags_by_row[i].append(_flag("subject", f"unknown subject '{subject}' for grade {grade}"))
        # NOTE: deliberately NOT flagging multiple teachers assigned to the
        # same (subject, grade, section) — golden rule 2, that's valid.

    # A teacher with no subject anywhere in the file is not necessarily a
    # problem: the usual meaning is "takes everything for that class". The row
    # gets every subject configured for its grade (see buildStudentsPlan's twin
    # in useImport.js buildTeachersPlan), so the flag says what WILL happen and
    # asks for verification rather than reporting an absence.
    #
    # It stays a flag, not silence, because a school with subject specialists —
    # common from Grade VI up — must be able to find and correct these.
    for tkey, idxs in by_teacher.items():
        if not all(not rows[i].get("subject", "").strip() for i in idxs):
            continue
        for i in idxs:
            grade = normalize_grade(rows[i].get("grade"))
            available = subject_names_by_grade.get(grade)
            if available:
                flags_by_row[i].append(_flag(
                    "subject",
                    f"subjects inferred from class assignment — verify "
                    f"({len(available)} subject(s) configured for grade {grade})"))
            else:
                flags_by_row[i].append(_flag(
                    "subject",
                    f"no subjects configured for grade {grade or '?'} — "
                    "assign manually after configuring the Subjects tab"))

    return flags_by_row

def validate_required_fields(entity, rows):
    required = SCHEMAS[entity]["required"]
    flags_by_row = []
    for r in rows:
        flags = []
        for k in required:
            if not r.get(k, "").strip():
                flags.append(_flag(k, f"missing {k.replace('_', ' ')}", severity="error"))
        flags_by_row.append(flags)
    return flags_by_row

def core_subject_coverage_flags(rows, class_lookup, subjects_by_class, core_subjects_by_grade,
                                class_by_doc_id=None):
    """Job-level (not per-row) check: any (class, core subject) with zero
    teacher rows covering it. Best-effort — 'core' comes from live subjects'
    `area` field when present, else falls back to every subject already
    attached to the class in `classes/{id}.subjects[]` (no area data to
    narrow it further)."""
    covered = set()
    for r in rows:
        gkey = (normalize_grade(r.get("grade")), normalize_section(r.get("section")))
        classId = class_lookup.get(gkey) or resolve_by_doc_id(class_by_doc_id, r.get("grade"), r.get("section"))
        subject = r.get("subject", "").strip().rstrip("?")
        if classId and subject:
            covered.add((classId, subject.lower()))

    out = []
    for classId, subject_ids in subjects_by_class.items():
        grade = classId.split("_")[0] if "_" in classId else classId
        core = core_subjects_by_grade.get(grade)
        candidates = core if core is not None else subject_ids
        for subjectId in candidates:
            subj_name = subjectId.split("_", 1)[-1].replace("_", " ")
            if (classId, subj_name.lower()) not in covered and (classId, subjectId.lower()) not in covered:
                out.append(f"{classId}: no teacher assigned for core subject '{subjectId}'")
    return out

# --------------------------------------------------------- school config ----
def load_school_config(db, school_id, entity):
    """Returns a dict bundle (not a tuple — too many fields now that cleaning
    needs canonical-keyed lookups alongside the original membership checks):
      class_lookup            (grade, section) -> classId, unchanged
      class_by_doc_id         class_id_key(doc id) -> classId, for the rows
                               whose file names the class by its whole id
                               rather than as a Class + Section pair
                               (see resolve_by_doc_id)
      subjects_by_class       classId -> [subjectId, ...], unchanged
      sections_by_grade       grade -> {canonical_section: display_section},
                               for match_value() in clean_students_rows/
                               clean_teachers_rows, and to tell "grade not
                               configured at all" from "grade configured but
                               not this section" in validate_students/teachers
      subject_names_by_grade  grade -> {lowercased name, ...}, unchanged
                               (validate_teachers' existing membership check)
      subjects_by_grade_canon grade -> {canonical_subject: display_subject},
                               for match_value() in clean_teachers_rows
      core_subjects_by_grade  unchanged
      aliases                 global import_aliases map, see load_import_aliases
      kb_overlay              learned education-KB overlay, see load_kb_entries
    """
    classes_ref = db.collection("schools").document(school_id).collection("classes")
    classes = [{"id": d.id, **d.to_dict()} for d in classes_ref.stream()]
    class_lookup = {}
    class_by_doc_id = {}
    subjects_by_class = {}
    sections_by_grade = {}
    for c in classes:
        grade = normalize_grade(c.get("clazz"))
        section_display = (c.get("section") or "").strip()
        section_norm = normalize_section(section_display)
        class_lookup[(grade, section_norm)] = c["id"]
        class_by_doc_id[class_id_key(c["id"])] = c["id"]
        subjects_by_class[c["id"]] = [s.get("subjectId") for s in (c.get("subjects") or []) if s.get("subjectId")]
        if section_display:
            sections_by_grade.setdefault(grade, {})[canonicalize(section_display)] = section_norm

    subject_names_by_grade = None
    subjects_by_grade_canon = {}
    core_subjects_by_grade = {}
    if entity == "teachers":
        subjects_ref = db.collection("schools").document(school_id).collection("subjects")
        subject_docs = [{"id": d.id, **d.to_dict()} for d in subjects_ref.stream()]
        subject_names_by_grade = {}
        has_area = any("area" in s for s in subject_docs)
        for s in subject_docs:
            grade = normalize_grade(s["id"].split("_", 1)[0]) if "_" in s["id"] else ""
            name = (s.get("name") or "").strip()
            subject_names_by_grade.setdefault(grade, set()).add(name.lower())
            if name:
                subjects_by_grade_canon.setdefault(grade, {})[canonicalize(name)] = name
            if has_area and (s.get("area") or "").strip().lower() == "scholastic":
                core_subjects_by_grade.setdefault(grade, set()).add(s["id"])

    return {
        "class_lookup": class_lookup,
        "class_by_doc_id": class_by_doc_id,
        "subjects_by_class": subjects_by_class,
        "sections_by_grade": sections_by_grade,
        "subject_names_by_grade": subject_names_by_grade,
        "subjects_by_grade_canon": subjects_by_grade_canon,
        "core_subjects_by_grade": core_subjects_by_grade,
        "aliases": load_import_aliases(db),
        "kb_overlay": load_kb_entries(db),
    }

# ----------------------------------------------------- per-file extraction -
def _process_one_file(entity, name, raw, job_diag):
    """Returns (rows, provenance) for one file, aligned 1:1. Deterministic
    tabular parsing (tabular_parser.parse_tabular_file, content-sniffed —
    never trusts the extension) is tried first for the `students` entity;
    every other entity, and students when the deterministic parser can't
    find a usable header, goes through the existing LLM extraction path.
    Skips the LLM fallback on a genuinely terminal condition (empty/
    password-protected/corrupt file) — that's pointless and slow to retry,
    and would just bury the specific reason under a vaguer LLM error.
    job_diag accumulates cross-file diagnostics for the job doc:
    {warnings, errors, file_summaries, unmapped_headers}."""
    if entity != "students":
        llm_rows = extract_file(entity, name, raw)
        provenance = [{"source_file": name, "sheet": None, "row": None,
                        "excluded_default": False, "extra_flags": []} for _ in llm_rows]
        job_diag["file_summaries"].append({
            "name": name, "file_type_detected": None, "sheets_parsed": [], "sheets_skipped": [],
            "row_count": len(llm_rows), "warning_count": 0, "error_count": 0, "used_llm_fallback": True,
        })
        return llm_rows, provenance

    tab = parse_tabular_file(name, raw, STUDENT_SCHEMA_KEYS, STUDENT_HEADER_ALIASES, STUDENT_REQUIRED_FIELD)
    job_diag["unmapped_headers"].update(tab.get("unmapped_headers", []))
    terminal = any(e["message"].startswith(("empty file", "file unreadable")) for e in tab["errors"])

    rows, provenance = [], []
    used_fallback = False
    if tab["rows"]:
        warn_by_key, err_by_key = {}, {}
        for w in tab["warnings"]:
            if w.get("row") is not None:
                warn_by_key.setdefault((w["sheet"], w["row"]), []).append(w)
        for e in tab["errors"]:
            if e.get("row") is not None:
                err_by_key.setdefault((e["sheet"], e["row"]), []).append(e)
        for r in tab["rows"]:
            rows.append(r["data"])
            key = (r["sheet"], r["row"])
            extra_flags = ([_flag(w["field"], w["message"], severity=w.get("severity", "warning"))
                             for w in warn_by_key.get(key, [])]
                            + [_flag(e["field"], e["message"], severity=e.get("severity", "error"))
                               for e in err_by_key.get(key, [])])
            extras = {k: v for k, v in (r.get("extras") or {}).items()
                      if not is_banned_extra(k)}
            dropped = sorted(k for k in (r.get("extras") or {}) if is_banned_extra(k))
            if dropped:
                extra_flags = extra_flags + [_flag(
                    None, f"not stored (golden rule 3): {', '.join(dropped)}")]
            provenance.append({
                "source_file": name, "sheet": r["sheet"], "row": r["row"],
                "excluded_default": bool(r.get("excluded")), "extra_flags": extra_flags,
                "extras": extras,
            })
    elif not terminal:
        used_fallback = True
        try:
            llm_rows = extract_file(entity, name, raw)
        except Exception as e:
            llm_rows = []
            job_diag["errors"].append({"file": name, "sheet": None, "row": None, "field": None,
                                        "severity": "error", "message": f"extraction failed: {e}"})
        for r in llm_rows:
            rows.append(r)
            provenance.append({"source_file": name, "sheet": None, "row": None,
                                "excluded_default": False, "extra_flags": []})

    for w in tab["warnings"]:
        job_diag["warnings"].append({"file": name, "sheet": w["sheet"], "row": w["row"],
                                      "field": w["field"], "severity": w.get("severity", "warning"),
                                      "message": w["message"]})
    for e in tab["errors"]:
        job_diag["errors"].append({"file": name, "sheet": e["sheet"], "row": e["row"],
                                    "field": e["field"], "severity": e.get("severity", "error"),
                                    "message": e["message"]})
    job_diag["file_summaries"].append({
        "name": name, "file_type_detected": tab["file_type_detected"],
        "sheets_parsed": tab["sheets_parsed"], "sheets_skipped": tab["sheets_skipped"],
        "row_count": len(rows), "warning_count": len(tab["warnings"]), "error_count": len(tab["errors"]),
        "used_llm_fallback": used_fallback,
    })
    return rows, provenance


def _log_unknown_headers(db, labels, school_id):
    """Firestore `import_unknown_headers`: {label, count, last_seen,
    last_school_id} keyed by canonicalized label, so the alias dictionary in
    normalize.STUDENT_HEADER_ALIASES can be grown from real production data
    instead of guesswork — see the task's NOTES section."""
    for label in labels:
        canon = canonicalize(label)
        if not canon:
            continue
        try:
            db.collection("import_unknown_headers").document(canon).set({
                "label": label, "count": firestore.Increment(1),
                "last_seen": firestore.SERVER_TIMESTAMP, "last_school_id": school_id,
            }, merge=True)
        except Exception as e:
            print(f"_log_unknown_headers: could not log '{label}': {e}")


# ------------------------------------------------------------------ main ----
@https_fn.on_call(
    region="asia-south1",
    memory=options.MemoryOption.GB_1,
    timeout_sec=540,
    max_instances=3,
    # OPENAI_API_KEY is the active provider (see extract_file's provider
    # switch above); ANTHROPIC_API_KEY still works if bound instead — e.g.
    # for .pdf uploads, which the OpenAI path can't extract from directly.
    secrets=["OPENAI_API_KEY"],
)
def process_import(req: https_fn.CallableRequest):
    _require_ops_admin(req)

    data = req.data or {}
    # Logged unconditionally (not just on failure) so a payload-shape
    # regression on the frontend shows up in Cloud Functions logs even when
    # the callable framework itself rejects the request before this line
    # (e.g. a stale build still POSTing an unwrapped body — the framework's
    # own pre-execution check raises INVALID_ARGUMENT "Bad Request" for that
    # case, with no message of ours to log; req.data is at least visible for
    # any request that *does* reach here).
    print(f"process_import: received payload keys={sorted(data.keys())}")

    # Callable requests cap out around 10MB; files[] here only ever carries
    # Storage paths (uploadAndProcess in useImport.js sends raw bytes to
    # Storage first, never inline), so this should stay a few hundred bytes.
    # Logged as a guard against that invariant ever regressing.
    approx_size = len(json.dumps(data, default=str))
    if approx_size > 8 * 1024 * 1024:
        print(f"process_import: payload is {approx_size} bytes — approaching the ~10MB callable limit")

    school_id = (data.get("schoolId") or "").strip()
    job_id = (data.get("jobId") or "").strip()
    entity = (data.get("entity") or "").strip()
    files = data.get("files") or []

    missing = []
    if not school_id:
        missing.append("schoolId")
    if not job_id:
        missing.append("jobId")
    if entity not in SCHEMAS:
        missing.append(f"entity (got {entity!r}, expected one of {sorted(SCHEMAS)})")
    if not files:
        missing.append("files (missing or empty)")
    else:
        for i, f in enumerate(files):
            if not isinstance(f, dict) or not (f.get("path") or "").strip():
                missing.append(f"files[{i}].path")
    if missing:
        raise https_fn.HttpsError(
            https_fn.FunctionsErrorCode.INVALID_ARGUMENT,
            f"Missing/invalid field(s): {'; '.join(missing)}")

    # Everything below is inside the try — including client/bucket/job_ref
    # construction, which used to sit outside it and could throw an
    # exception the on_call framework catches on its own and redacts to a
    # bare INTERNAL error with no message. job_ref may still be None in the
    # except block if that construction itself is what failed.
    job_ref = None
    try:
        db = firestore.client()
        bucket = storage.bucket()
        job_ref = db.collection("staging_imports").document(job_id)

        job_ref.set({
            "school_id": school_id, "entity": entity,
            "source_files": [f["path"] for f in files],
            "status": "processing", "model_used": MODEL,
        }, merge=True)

        # job_diag accumulates cross-file diagnostics (never per-row data —
        # that stays on each row doc) for the job summary: which files
        # produced what, which sheets got skipped and why, every file/sheet/
        # row-scoped warning and error. Capped before being written to the
        # job doc (see below) to stay well under Firestore's 1MB doc limit
        # even for a very messy multi-hundred-row file.
        job_diag = {"warnings": [], "errors": [], "file_summaries": [], "unmapped_headers": set()}

        raw_rows = []
        provenance = []
        for f in files:
            name = f.get("name", f["path"])
            try:
                raw = bucket.blob(f["path"]).download_as_bytes()
                file_rows, file_prov = _process_one_file(entity, name, raw, job_diag)
            except Exception as e:
                # One bad file must never take down the whole job — degrade
                # to a file-level error and keep processing the rest.
                job_diag["errors"].append({"file": name, "sheet": None, "row": None, "field": None,
                                            "severity": "error", "message": f"could not process file: {e}"})
                job_diag["file_summaries"].append({
                    "name": name, "file_type_detected": "unknown", "sheets_parsed": [], "sheets_skipped": [],
                    "row_count": 0, "warning_count": 0, "error_count": 1, "used_llm_fallback": False,
                })
                continue
            raw_rows.extend(file_rows)
            provenance.extend(file_prov)

        if entity == "students" and job_diag["unmapped_headers"]:
            _log_unknown_headers(db, job_diag["unmapped_headers"], school_id)

        cfg = load_school_config(db, school_id, entity)

        # Cleaning (Stage 1 deterministic + Stage 2 alias/fuzzy) runs before
        # validation, on every parsed row — see the "cleaning" section above.
        # Teachers can come out LONGER than raw_rows: a multi-grade cell
        # ('Garde 7,10') expands into one row per grade. Students/teachers
        # match against the school's own config; subjects are classified
        # against the education knowledge base (scholastic vs co-scholastic);
        # assessments pass through unchanged (nothing in their schema needs
        # matching against anything).
        if entity == "students":
            cleaned = clean_students_rows(raw_rows, cfg, provenance)
        elif entity == "teachers":
            cleaned = clean_teachers_rows(raw_rows, cfg, provenance)
        elif entity == "subjects":
            cleaned = clean_subjects_rows(raw_rows, cfg, provenance)
        else:
            # No expansion for assessments (1 input row -> 1 output row), so
            # provenance stays index-aligned with raw_rows exactly like the
            # students path.
            cleaned = [{"data": r, "fixes": [], "suggestions": [], "flags": [],
                        "provenance": provenance[i] if i < len(provenance) else None}
                       for i, r in enumerate(raw_rows)]

        # Merge in tabular_parser's own per-row diagnostics (ambiguous dates,
        # invalid phone numbers, missing-name exclusions, ...) — only
        # meaningful for students, and only 1:1-aligned there (teachers can
        # expand past provenance's original length via multi-grade cells).
        if entity == "students":
            for c in cleaned:
                extra = (c.get("provenance") or {}).get("extra_flags") or []
                if extra:
                    c["flags"] = extra + c["flags"]
            drop_unrecognized_grade_flags(cleaned, cfg["class_by_doc_id"])

        cleaned_data = [c["data"] for c in cleaned]
        if entity == "students":
            validation_flags = validate_students(
                cleaned_data, cfg["class_lookup"], cfg["sections_by_grade"],
                existing_flags=[c["flags"] for c in cleaned],
                class_by_doc_id=cfg["class_by_doc_id"])
            class_level_flags = []
        elif entity == "teachers":
            validation_flags = validate_teachers(
                cleaned_data, cfg["class_lookup"], cfg["sections_by_grade"], cfg["subject_names_by_grade"],
                class_by_doc_id=cfg["class_by_doc_id"])
            class_level_flags = core_subject_coverage_flags(
                cleaned_data, cfg["class_lookup"], cfg["subjects_by_class"], cfg["core_subjects_by_grade"],
                class_by_doc_id=cfg["class_by_doc_id"])
        else:
            validation_flags = validate_required_fields(entity, cleaned_data)
            class_level_flags = []

        for c, v_flags in zip(cleaned, validation_flags):
            c["flags"] = c["flags"] + v_flags

        rows_ref = job_ref.collection("rows")
        flag_count = fixed_count = suggestion_count = error_flag_count = 0
        included_row_count = 0
        for i in range(0, len(cleaned), 450):
            batch = db.batch()
            for j in range(i, min(i + 450, len(cleaned))):
                c = cleaned[j]
                prov = c.get("provenance") or {}
                flag_count += len(c["flags"])
                error_flag_count += sum(1 for fl in c["flags"] if fl.get("severity") == "error")
                if c["fixes"]:
                    fixed_count += 1
                if c["suggestions"]:
                    suggestion_count += 1
                excluded_default = bool(prov.get("excluded_default"))
                if not excluded_default:
                    included_row_count += 1
                batch.set(rows_ref.document(str(j)), {
                    "data": c["data"], "extras": prov.get("extras") or {},
                    "flags": c["flags"],
                    "fixes": c["fixes"], "suggestions": c["suggestions"],
                    "edited": False, "excluded": excluded_default,
                    "source_file": prov.get("source_file"), "source_sheet": prov.get("sheet"),
                    "source_row": prov.get("row"),
                })
            batch.commit()

        # 0 (usable) rows is ALWAYS an error state, per the parser-hardening
        # contract — never "ready" with an empty table. A job can produce
        # rows that are ALL excluded (e.g. every row in the file was missing
        # a name) and that's the same story: nothing committable came out of
        # this job, so the UI must show an error, not a quiet empty table.
        status = "ready" if included_row_count > 0 else "error"
        if status == "error" and not job_diag["errors"]:
            job_diag["errors"].append({
                "file": None, "sheet": None, "row": None, "field": None, "severity": "error",
                "message": "0 usable rows extracted from the uploaded file(s)",
            })

        job_ref.set({
            "status": status, "row_count": len(cleaned), "included_row_count": included_row_count,
            "flag_count": flag_count, "error_flag_count": error_flag_count,
            "fixed_count": fixed_count, "suggestion_count": suggestion_count,
            "class_level_flags": class_level_flags,
            "file_summaries": job_diag["file_summaries"],
            "parse_warnings": job_diag["warnings"][:500],
            "parse_errors": job_diag["errors"][:500],
            # Source columns the parser saw but could not map to any field.
            # Previously only logged to import_unknown_headers (invisible to
            # the operator), so a column the file carried simply vanished from
            # Review with nothing said. Now surfaced on the job so the review
            # screen can name every dropped column.
            "unmapped_headers": sorted(job_diag["unmapped_headers"])[:100],
            "review_only_fields": REVIEW_ONLY_STUDENT_KEYS if entity == "students" else [],
            "completed_at": firestore.SERVER_TIMESTAMP,
        }, merge=True)

        return {"status": status, "rowCount": len(cleaned), "includedRowCount": included_row_count,
                "flagCount": flag_count, "fixedCount": fixed_count, "suggestionCount": suggestion_count}

    except https_fn.HttpsError:
        raise
    except Exception as e:
        if job_ref is not None:
            job_ref.set({"status": "failed", "error": str(e)}, merge=True)
        raise https_fn.HttpsError(https_fn.FunctionsErrorCode.INTERNAL, str(e))


# ------------------------------------------------------- classify_value -----
# LAST RESORT ONLY. The deterministic knowledge base (education_kb.py) answers
# the overwhelming majority of values for free; this callable exists purely
# for the ones it has never seen, and the frontend must only reach it after
# classify() has come back `unknown` (see src/composables/useEducationKB.js).
#
# Three hard rules, all enforced here rather than trusted to the caller:
#   - It is asked ONCE per value: the answer is cached in `kb_entries` on
#     confirmation, and a cached entry short-circuits classify() forever after.
#   - It NEVER writes to the KB. It returns a SUGGESTION. Only a human
#     accepting that suggestion writes the entry (client-side, recording who
#     confirmed it), so a hallucinated answer can't silently become canon.
#   - A refusal/garbage/unavailable model degrades to type "unknown" with a
#     reason, never to a guess.
CLASSIFY_PROMPT = (
    "You are classifying a single value taken from an Indian school's own data "
    "files, for a report-card platform.\n"
    "Value: {value}\n"
    "{context}"
    "Answer with ONE of these types:\n"
    '  "subject"       — a scholastic/academic subject (English, Mathematics, Science...)\n'
    '  "coscholastic"  — a co-scholastic area or activity (Art & Craft, Yoga, Physical Education...)\n'
    '  "grade"         — a class/grade level (Nursery, LKG, 1..12)\n'
    '  "section"       — a section/division name within a grade (A, Rose, Champion...)\n'
    '  "other"         — none of the above (a header label, a teacher name, a stray note...)\n'
    'Also give the canonical form you would file it under (for "other", use "").\n'
    "Return ONLY JSON: {{\"type\": \"...\", \"canonical\": \"...\", \"reason\": \"one short sentence\"}}"
)


def _llm_classify(value, context=""):
    """Returns {type, canonical, reason} or None when no model is reachable."""
    prompt = CLASSIFY_PROMPT.format(value=value, context=context)
    blocks = [{"type": "text", "text": prompt}]
    caller = call_openai_compatible if os.environ.get("OPENAI_API_KEY") else call_anthropic
    raw = caller(blocks, "Classify the value above.")
    raw = re.sub(r"^```(json)?|```$", "", (raw or "").strip(), flags=re.M).strip()
    data = json.loads(raw)
    kind = str(data.get("type") or "").strip().lower()
    if kind not in kb.ENTITY_TYPES + (kb.OTHER,):
        return None
    return {
        "type": kind,
        "canonical": str(data.get("canonical") or "").strip(),
        "reason": str(data.get("reason") or "").strip()[:300],
    }


@https_fn.on_call(region="asia-south1", memory=options.MemoryOption.MB_512,
                   timeout_sec=60, max_instances=3, secrets=["OPENAI_API_KEY"])
def classify_value(req: https_fn.CallableRequest):
    """Classify one unrecognized value. Request: {value, context?}.

    Response: {source, type, canonical, confidence, reason}. `source` is
    "kb" when the deterministic layer actually knew the answer after all
    (the client raced a freshly-learned entry, or passed a value it hadn't
    checked), "llm" for a model suggestion, "unavailable" when no model
    answered. NOTHING here writes to Firestore.
    """
    _require_ops_admin(req)

    data = req.data or {}
    value = (data.get("value") or "").strip()
    context = (data.get("context") or "").strip()
    if not value:
        raise https_fn.HttpsError(
            https_fn.FunctionsErrorCode.INVALID_ARGUMENT, "Missing value")

    try:
        db = firestore.client()
        overlay = load_kb_entries(db)
    except Exception as e:
        print(f"classify_value: KB overlay unavailable, using seed only: {e}")
        overlay = {}

    # Never spend a model call on something already known — the client is
    # supposed to check first, but this is the guarantee, not the hope.
    known = kb.classify(value, overlay=overlay)
    if known["type"] != kb.UNKNOWN and known["confidence"] >= kb.MEDIUM_CONFIDENCE:
        return {"source": "kb", "type": known["type"], "canonical": known["canonical"],
                "confidence": known["confidence"],
                "reason": f"Already known to the knowledge base ({known['source']} match)."}

    ctx = f"Context: it appeared in {context}.\n" if context else ""
    try:
        suggestion = _llm_classify(value, ctx)
    except Exception as e:
        print(f"classify_value: LLM call failed for {value!r}: {e}")
        suggestion = None

    if not suggestion:
        return {"source": "unavailable", "type": kb.UNKNOWN, "canonical": "",
                "confidence": 0.0,
                "reason": "No classification available — set the type manually."}

    return {"source": "llm", "type": suggestion["type"],
            "canonical": suggestion["canonical"] or value,
            # Deliberately below HIGH_CONFIDENCE: a model answer is a
            # suggestion awaiting human confirmation, never an auto-apply.
            "confidence": 0.5, "reason": suggestion["reason"]}


# --------------------------------------------------------------- commit -----
# Takes the plan the review UI already built client-side (useImport.js's
# buildCommitPlan — plain Firestore reads against classes/subjects/staffs,
# unchanged) and performs the actual writes into schools/{schoolId}/... here,
# server-side, so the write path gets the same req.auth + allowlist check as
# extraction instead of relying on the client's own Firestore permissions.
# `items` mirrors plan.items from useImport.js, trimmed to just what's needed
# to write: {status, docId, payload} for students/subjects/assessments, or
# {status, staffId, staffIsNew, classId, subjectId, staffBase} for teachers.
@https_fn.on_call(region="asia-south1", memory=options.MemoryOption.MB_512,
                   timeout_sec=120, max_instances=3)
def commit_import(req: https_fn.CallableRequest):
    email = _require_ops_admin(req)

    data = req.data or {}
    school_id = (data.get("schoolId") or "").strip()
    job_id = (data.get("jobId") or "").strip()
    entity = (data.get("entity") or "").strip()
    items = data.get("items") or []
    overwrite_existing = bool(data.get("overwriteExisting"))

    if not school_id or not job_id or entity not in ("students", "teachers", "subjects", "assessments"):
        raise https_fn.HttpsError(
            https_fn.FunctionsErrorCode.INVALID_ARGUMENT,
            "Missing schoolId, jobId, or valid entity")

    # Everything below is inside the try — including client/ref construction,
    # which used to sit outside it and could throw an exception the on_call
    # framework catches on its own and redacts to a bare INTERNAL error with
    # no message (see process_import's identical fix above).
    try:
        db = firestore.client()
        school_ref = db.collection("schools").document(school_id)
        job_ref = db.collection("staging_imports").document(job_id)

        writable = [it for it in items if it.get("status") == "CREATE"
                    or (it.get("status") == "UPDATE_CHANGED" and overwrite_existing)]

        # THE GATE. This function writes with the Admin SDK, which bypasses
        # Firestore security rules entirely, and _commit_simple used to write
        # the browser's payload verbatim. The dashboard validates too, but a
        # stale tab or a hand-made call would sail straight past that; nothing
        # malformed gets into a school's tree from here.
        #
        # Per-row, not all-or-nothing: valid rows commit and rejected ones come
        # back with a reason, matching how the import preview already behaves.
        writable, rejected = _validate_writable(entity, writable)

        if entity in ("students", "subjects"):
            _commit_simple(db, school_ref.collection(entity), writable, email)
        elif entity == "assessments":
            _commit_assessments(db, school_ref.collection("assessments"), writable, email)
        elif entity == "teachers":
            _commit_teachers(school_ref.collection("staffs"), writable, email)

        job_ref.set({
            "status": "committed", "committed_at": firestore.SERVER_TIMESTAMP,
            "committed_by": email,
            "rejected": rejected,
        }, merge=True)

        return {"written": len(writable), "skipped": len(items) - len(writable),
                "rejected": rejected}

    except https_fn.HttpsError:
        raise
    except Exception as e:
        raise https_fn.HttpsError(https_fn.FunctionsErrorCode.INTERNAL, str(e))


# The import entity name and the collection it lands in are not always the
# same word — "teachers" writes into staffs, "assessments" into assessments.
_ENTITY_COLLECTION = {
    "students": "students", "subjects": "subjects",
    "assessments": "assessments", "teachers": "staffs",
}


def _validate_writable(entity, writable):
    """Split rows into (accepted, rejected) against the shared schema.

    `teachers` is skipped: _commit_teachers builds its own staff document from
    assignment fields rather than carrying a payload, so there is nothing here
    to validate against a doc shape. It is gated by the same ops-admin check
    as everything else.
    """
    collection = _ENTITY_COLLECTION.get(entity)
    if not collection or entity == "teachers":
        return writable, []

    accepted, rejected = [], []
    for item in writable:
        doc_id = item.get("docId") or ""
        payload = coerce_wire_payload(collection, item.get("payload") or {})

        if collection == "students":
            cls = validate_current_class_id(payload.get("currentClassId"), student_id=doc_id)
            if not cls["ok"]:
                rejected.append({"docId": doc_id, "reason": cls["message"]})
                continue

        result = validate_doc(collection, payload)
        if not result["ok"]:
            rejected.append({"docId": doc_id,
                             "reason": f"does not match the {collection} schema — "
                                       f"{format_errors(result['errors'])}"})
            continue

        # Carry the coerced payload forward — the ISO string the browser sent
        # is now a real datetime, which is what must reach Firestore.
        accepted.append({**item, "payload": payload})

    return accepted, rejected


def _commit_simple(db, collection_ref, writable, email):
    """students / subjects: one doc per item, merge-set."""
    for i in range(0, len(writable), 450):
        batch = db.batch()
        for item in writable[i:i + 450]:
            payload = dict(item.get("payload") or {})
            payload["updated_at"] = firestore.SERVER_TIMESTAMP
            payload["updated_by"] = email
            if item.get("status") == "CREATE":
                payload["created_at"] = firestore.SERVER_TIMESTAMP
                payload["created_by"] = email
            batch.set(collection_ref.document(item["docId"]), payload, merge=True)
        batch.commit()


def _commit_assessments(db, collection_ref, writable, email):
    """Assessments are always CREATE (buildAssessmentsPlan never emits
    UPDATE_* — see its comment in useImport.js), but always stamp both
    created_* and updated_* to match the original client-side behavior."""
    for i in range(0, len(writable), 450):
        batch = db.batch()
        for item in writable[i:i + 450]:
            payload = dict(item.get("payload") or {})
            payload["updated_at"] = firestore.SERVER_TIMESTAMP
            payload["updated_by"] = email
            payload["created_at"] = firestore.SERVER_TIMESTAMP
            payload["created_by"] = email
            batch.set(collection_ref.document(item["docId"]), payload, merge=True)
        batch.commit()


def _commit_teachers(staffs_ref, writable, email):
    """Group by staff so each teacher gets exactly one write with a merged
    assignments map — never overwrite another subject's assignment for the
    same class, and never drop a second teacher on the same
    (subject, grade, section): that's staffs/{idA} and staffs/{idB} each
    separately gaining the same classId+subjectId in their own doc."""
    by_staff = {}
    for item in writable:
        staff_id = item.get("staffId")
        if staff_id not in by_staff:
            by_staff[staff_id] = {
                "staffIsNew": item.get("staffIsNew"),
                "staffBase": item.get("staffBase") or {},
                "adds": [],
            }
        by_staff[staff_id]["adds"].append(item)

    for staff_id, group in by_staff.items():
        assignments = {}
        class_ids = []
        if not group["staffIsNew"]:
            snap = staffs_ref.document(staff_id).get()
            if snap.exists:
                d = snap.to_dict() or {}
                assignments = dict(d.get("assignments") or {})
                class_ids = list(d.get("classIds") or [])
        class_id_set = set(class_ids)
        for item in group["adds"]:
            # subjectIds is the list the planner builds — one entry for an
            # explicit subject cell, every subject of the grade when the cell
            # was blank and the default was inferred. subjectId remains the
            # fallback so a stale client (or a replayed job) still commits.
            subject_ids = [s for s in (item.get("subjectIds") or []) if s]
            if not subject_ids and item.get("subjectId"):
                subject_ids = [item["subjectId"]]
            class_id = item.get("classId")
            if not class_id:
                continue
            # A class with no subjects still grants class-level access
            # (attendance, remarks, co-scholastic all key off classIds), so the
            # class is recorded even when the subject list is empty.
            class_id_set.add(class_id)
            if not subject_ids:
                continue
            existing = set(assignments.get(class_id) or [])
            existing.update(subject_ids)
            assignments[class_id] = sorted(existing)

        payload = {
            "assignments": assignments, "classIds": list(class_id_set),
            "updated_at": firestore.SERVER_TIMESTAMP, "updated_by": email,
        }
        if group["staffIsNew"]:
            base = group["staffBase"]
            payload.update({
                "id": staff_id, "name": base.get("name") or "", "email": base.get("email") or "",
                "type": "teacher", "needsAuthCreation": True, "authUid": None,
                "created_at": firestore.SERVER_TIMESTAMP, "created_by": email,
            })
        staffs_ref.document(staff_id).set(payload, merge=True)
