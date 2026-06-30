"""
ClarifiEd Agreement Generator — Cloud Function
Fills variables into the agreement Word template and returns:
  - JSON with base64-encoded PDF and DOCX

Deploy:
  gcloud functions deploy generate_agreement \
    --gen2 --runtime python312 --region asia-south1 \
    --source . --entry-point generate_agreement \
    --trigger-http --allow-unauthenticated \
    --memory 512MB --max-instances 3 --project clarified-1501

Folder needs: main.py, requirements.txt, agreement_template.docx
LibreOffice must be available for PDF conversion.
"""

import base64
import copy
import io
import json
import os
import re
import subprocess
import tempfile
from datetime import datetime

import functions_framework
from docx import Document
from docx.shared import Pt
from flask import Request, Response

API_KEY = "9421060748"
_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(_DIR, "agreement_template.docx")


def _current_academic_year():
    """India academic year: April-March. e.g. Jan-Mar 2027 -> '2026-27', Apr 2026 onwards -> '2026-27'"""
    today = datetime.today()
    if today.month >= 4:
        start = today.year
    else:
        start = today.year - 1
    return f"{start}-{str(start + 1)[-2:]}"


def _format_inr(amount):
    s = str(int(amount))
    if len(s) <= 3:
        return s
    result = s[-3:]
    s = s[:-3]
    while s:
        result = s[-2:] + "," + result
        s = s[:-2]
    return result.lstrip(",")


def _get_installment_splits(plan):
    return [50, 25, 25] if plan != "B" else [25, 25, 25, 25]


def _fix_textbox_date(doc):
    """
    The signature date ('Date: 11/10/2025') lives inside a Word text box,
    which python-docx's doc.paragraphs does not traverse. We find and
    blank out the static date here so it reads just 'Date: ' for the
    School to fill in by hand (this Agreement is generated per-school,
    not signed digitally, so a blank date field is correct).
    """
    from docx.oxml.ns import qn
    body = doc.element.body
    # Find all w:t elements anywhere in the document (covers text boxes)
    for t in body.iter(qn('w:t')):
        if t.text and '11/10/2025' in t.text:
            t.text = t.text.replace('11/10/2025', '')
        elif t.text == '11/10/2025':
            t.text = ''


def _replace_in_para(para, old, new):
    """
    Replace text in a paragraph that may be split across multiple runs.
    Strategy: if the exact 'old' substring exists in the concatenated
    paragraph text, rebuild the paragraph text with the replacement,
    placing all of it in the first run and clearing the rest. This loses
    fine-grained character formatting within the replaced span but
    preserves the paragraph-level formatting (bold headers stay bold etc),
    which is what we need for legal docs like this one.
    """
    full = "".join(run.text for run in para.runs)
    if old not in full:
        return False

    new_full = full.replace(old, new)

    if not para.runs:
        return False

    # Put all new text in first run, preserving its formatting
    para.runs[0].text = new_full
    # Clear all other runs
    for run in para.runs[1:]:
        run.text = ""
    return True


def _fill_agreement(data):
    school_name    = data["schoolName"]
    school_address = data.get("schoolAddress", "")
    signatory_name = data.get("signatoryName", "")
    signatory_desig= data.get("signatoryDesignation", "")
    hpc_type       = data.get("hpcType", "printed and digital")
    fee            = int(data["feePerStudent"])
    student_count  = int(data["studentCount"])
    plan           = data.get("installmentPlan", "A")
    agreement_num  = data.get("agreementNumber", "")

    total = fee * student_count
    splits = _get_installment_splits(plan)
    academic_year = _current_academic_year()

    # Build payment terms string
    if plan == "A":
        payment_terms = f"50% payable upon onboarding, 25% after Term 1 delivery, and 25% after final delivery"
    else:
        payment_terms = f"25% payable upon onboarding, 25% after Term 1, 25% after Term 2, and 25% after final delivery"

    # Build installment breakdown
    inst_lines = []
    for i, pct in enumerate(splits):
        amt = round(total * pct / 100)
        inst_lines.append(f"Instalment {i+1}: {pct}% — Rs. {_format_inr(amt)}/-")
    inst_str = "; ".join(inst_lines)

    doc = Document(TEMPLATE_PATH)

    # Define replacements: (old_text, new_text)
    # IMPORTANT: these must match the exact concatenated paragraph text
    # in agreement_template.docx (school name, address, fee, terms)
    replacements = [
        # School party intro — exact match from template para 3
        (
            "Samarth Dnyanpeeth's Sahyadri Public School, an educational institution located at Waghalwadi, Soratewadi, Maharashtra 412306 and represented herein by its authorised signatory (hereinafter referred to as the \u201cSchool,\u201d which expression shall, unless repugnant to the context or meaning thereof, include its successors and permitted assigns).",
            f"{school_name}, an educational institution located at {school_address} and represented herein by its authorised signatory (hereinafter referred to as the \u201cSchool,\u201d which expression shall, unless repugnant to the context or meaning thereof, include its successors and permitted assigns)."
        ),
        # Commercial terms paragraph — exact match from template para 19
        (
            "Consideration & Payment Terms: (a) The fee per student will be Rs. 131/- (including all taxes). The School shall pay the Company the agreed consideration, with 50% payable upon onboarding, and the remaining 50% after the final delivery of all soft and hard copies of HPCs. (b) The full fee shall be applicable for mid-year admissions. (c) All payments shall be due and payable within forty-five (45) working days of invoice and shall be non-refundable unless otherwise agreed.",
            f"Consideration & Payment Terms: (a) The fee per student will be Rs. {fee}/- (including all taxes). The School shall pay the Company the agreed consideration, with {payment_terms}. (b) The full fee shall be applicable for mid-year admissions. (c) All payments shall be due and payable within forty-five (45) working days of invoice and shall be non-refundable unless otherwise agreed. Total contract value based on {student_count} students: Rs. {_format_inr(total)}/-. Instalment breakdown: {inst_str}."
        ),
        # Signature school name — exact match from template para 81
        ("For Samarth Dnyanpeeth's Sahyadri Public School :", f"For {school_name} :"),
        # Remove date range — exact match including trailing space from template
        ("This Agreement shall be valid and effective from 1st June 2025  to                      31st  May  2026 ",
         f"This Agreement shall be valid and effective for the current academic year {academic_year}."),
        # HPC delivery type — template says "hard copy", not "printed and digital"
        ("the hard copy HPC within six (6) working days", f"the {hpc_type} HPC within six (6) working days"),
    ]

    for para in doc.paragraphs:
        for old, new in replacements:
            _replace_in_para(para, old, new)

    _fix_textbox_date(doc)

    # Note: agreement reference number is used in filename and can be
    # added to footer separately if needed — skipping inline insertion
    # into the title area since it disrupts document flow.

    return doc


def _docx_to_bytes(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _docx_bytes_to_pdf(docx_bytes):
    """Convert docx bytes to PDF using LibreOffice headless."""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "agreement.docx")
        pdf_path  = os.path.join(tmpdir, "agreement.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", tmpdir, docx_path],
            capture_output=True, text=True, timeout=60
        )

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice failed: {result.stderr}")

        with open(pdf_path, "rb") as f:
            return f.read()


@functions_framework.http
def generate_agreement(request: Request):
    # CORS preflight
    if request.method == "OPTIONS":
        return Response("", 204, headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "POST",
            "Access-Control-Allow-Headers": "Content-Type, X-Api-Key",
        })

    if request.headers.get("X-Api-Key") != API_KEY:
        return Response("Unauthorized", 401, headers={"Access-Control-Allow-Origin": "*"})

    data = request.get_json(silent=True) or {}

    school_name = (data.get("schoolName") or "").strip()
    if not school_name:
        return Response("Missing schoolName", 400, headers={"Access-Control-Allow-Origin": "*"})
    if not data.get("feePerStudent"):
        return Response("Missing feePerStudent", 400, headers={"Access-Control-Allow-Origin": "*"})
    if not data.get("studentCount"):
        return Response("Missing studentCount", 400, headers={"Access-Control-Allow-Origin": "*"})

    try:
        doc = _fill_agreement(data)
        docx_bytes = _docx_to_bytes(doc)
        pdf_bytes  = _docx_bytes_to_pdf(docx_bytes)
    except Exception as e:
        return Response(
            json.dumps({"error": str(e)}), 500,
            mimetype="application/json",
            headers={"Access-Control-Allow-Origin": "*"}
        )

    safe_name = "".join(ch for ch in school_name if ch not in '\\/:*?"<>|')

    response_data = {
        "docx": base64.b64encode(docx_bytes).decode(),
        "pdf":  base64.b64encode(pdf_bytes).decode(),
        "filename": f"Agreement_{safe_name}_{data.get('agreementNumber', '')}",
    }

    return Response(
        json.dumps(response_data),
        200,
        mimetype="application/json",
        headers={"Access-Control-Allow-Origin": "*"},
    )
